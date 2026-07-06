"""Generate comprehensive Word document for Viraasat project."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# ── Styles ──────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for level in range(1, 5):
    h = doc.styles[f"Heading {level}"]
    h.font.color.rgb = RGBColor(0x80, 0x00, 0x20)

# ── Title Page ──────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("VIRAASAT")
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x80, 0x00, 0x20)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Ethnic Women's Kurtis — E-Commerce Platform")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0xC5, 0xA0, 0x28)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(f"Project Report — {datetime.date.today().strftime('%B %Y')}")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x6B, 0x6B, 0x6B)

doc.add_page_break()

# ── Table of Contents ───────────────────────────────────
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Project Overview",
    "3. Technology Stack",
    "4. Architecture & Design",
    "5. Key Features",
    "6. User Guide",
    "7. Admin Guide",
    "8. Development & Deployment",
    "9. Known Limitations & Future Roadmap",
    "10. Appendix",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ── 1. Executive Summary ───────────────────────────────
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "Viraasat is a modern e-commerce platform built for selling premium ethnic women's "
    "kurtis online. The platform delivers a rich, interactive shopping experience with "
    "features including product browsing with advanced filtering, a wishlist and cart system, "
    "secure checkout, order tracking, customer reviews, and a full admin dashboard for "
    "managing products, orders, customers, and reviews."
)
doc.add_paragraph(
    "The project follows a modern two-tier architecture with a Next.js frontend and a "
    "FastAPI Python backend, using Supabase as the database and authentication provider. "
    "The application is designed to work fully offline in demo mode, making it easy to "
    "showcase and test without external dependencies."
)

# ── 2. Project Overview ────────────────────────────────
doc.add_heading("2. Project Overview", level=1)
doc.add_heading("2.1 Purpose", level=2)
doc.add_paragraph(
    "Viraasat aims to provide a seamless online shopping experience for ethnic wear, "
    "specifically targeting women seeking premium kurtis. The name 'Viraasat' (meaning "
    "'heritage' in Hindi) reflects the brand's focus on traditional craftsmanship combined "
    "with contemporary design."
)
doc.add_heading("2.2 Target Audience", level=2)
bullet = ["Fashion-conscious women aged 20–45", "Customers who prefer ethnic wear", "Online shoppers looking for premium kurtis"]
for item in bullet:
    doc.add_paragraph(item, style="List Bullet")
doc.add_heading("2.3 Core Business Goals", level=2)
goals = [
    "Provide an intuitive and visually appealing shopping interface",
    "Enable easy product discovery through categories and search",
    "Offer a smooth checkout process with address management",
    "Allow customers to track orders and manage their profiles",
    "Give administrators full control over products, orders, and content",
]
for g in goals:
    doc.add_paragraph(g, style="List Bullet")

# ── 3. Technology Stack ────────────────────────────────
doc.add_heading("3. Technology Stack", level=1)
doc.add_heading("3.1 Frontend", level=2)
frontend = [
    ("Framework", "Next.js 16 (Turbopack) — React-based framework with server-side rendering"),
    ("Language", "TypeScript (strict mode) — type-safe JavaScript"),
    ("Styling", "Tailwind CSS v4 — utility-first CSS framework"),
    ("State Management", "Zustand — lightweight state management for cart and wishlist"),
    ("Animation", "Motion (Framer Motion) — smooth page and UI animations"),
    ("UI Components", "Radix UI primitives + custom components"),
    ("Notifications", "Sonner — toast notification library"),
]
for label, value in frontend:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value)

doc.add_heading("3.2 Backend", level=2)
backend = [
    ("Framework", "FastAPI (Python) — high-performance async API server"),
    ("Database", "Supabase (PostgreSQL) — managed database with real-time capabilities"),
    ("Authentication", "Supabase Auth — email/password authentication with session management"),
    ("API Style", "RESTful JSON API with auto-generated OpenAPI docs"),
]
for label, value in backend:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(value)

doc.add_heading("3.3 Infrastructure", level=2)
infra = [
    "Frontend hosted as a static Next.js build or server-rendered app",
    "Backend served via Uvicorn ASGI server on port 8000",
    "Demo mode toggle via NEXT_PUBLIC_USE_DEMO=true environment variable",
    "Images served from Supabase storage (with local fallback)",
]
for item in infra:
    doc.add_paragraph(item, style="List Bullet")

# ── 4. Architecture & Design ───────────────────────────
doc.add_heading("4. Architecture & Design", level=1)
doc.add_heading("4.1 System Architecture", level=2)
doc.add_paragraph(
    "The application follows a classic two-tier architecture. The frontend (Next.js) "
    "communicates with the backend (FastAPI) via REST API calls. The backend interfaces "
    "with Supabase for data persistence and authentication. In demo mode, the frontend "
    "operates entirely offline using embedded mock data, eliminating the need for a "
    "running backend or Supabase instance."
)
doc.add_heading("4.2 Frontend Structure", level=2)
doc.add_paragraph(
    "The frontend is organized into the following directory structure:"
)
structure = [
    "app/ — Next.js App Router pages and layouts",
    "  (auth)/ — Login and signup pages",
    "  admin/ — Admin dashboard with sub-routes",
    "  cart/ — Shopping cart page",
    "  checkout/ — Checkout flow",
    "  orders/ — Order history and detail pages",
    "  products/ — Product listing and detail pages",
    "  profile/ — User profile and address management",
    "  wishlist/ — Wishlist page",
    "components/ — Reusable UI components",
    "lib/ — Business logic, stores, API clients, and demo data",
    "  supabase/ — Supabase client factories (client + server)",
    "  context/ — React context providers (auth)",
    "  stores/ — Zustand state stores",
]
for item in structure:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("4.3 Backend Structure", level=2)
backend_structure = [
    "app/main.py — FastAPI application entry point",
    "app/routes/ — API route modules (auth, admin, orders, etc.)",
    "app/core/ — Core infrastructure (config, Supabase client)",
    "app/models/ — Data models and types",
    "app/schemas/ — Pydantic request/response schemas",
    "app/services/ — Business logic layer",
]
for item in backend_structure:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("4.4 Data Flow", level=2)
doc.add_paragraph(
    "When a user browses products, the frontend queries either the Supabase database "
    "(live mode) or local demo data (demo mode). Cart and wishlist state is managed "
    "client-side via Zustand and persisted to Supabase when a user is logged in. "
    "Orders are placed through the FastAPI backend, which validates the request and "
    "records the order in Supabase. The admin dashboard fetches aggregated data from "
    "the backend API (live mode) or shows placeholder data (demo mode)."
)

# ── 5. Key Features ────────────────────────────────────
doc.add_heading("5. Key Features", level=1)
doc.add_heading("5.1 User Features", level=2)
user_features = [
    "Product browsing with category filtering and search",
    "Product detail pages with multiple images and size guide",
    "Add to cart with quantity management",
    "Wishlist for saving favorite products",
    "User authentication (sign up / log in)",
    "Checkout flow with address selection and order placement",
    "Order history and detailed order tracking",
    "Product reviews and ratings",
    "Responsive design for mobile and desktop",
    "Dark mode support",
]
for f in user_features:
    doc.add_paragraph(f, style="List Bullet")

doc.add_heading("5.2 Admin Features", level=2)
admin_features = [
    "Dashboard with key metrics (products, orders, revenue)",
    "Product management (add, edit, delete products)",
    "Order management with status updates",
    "Customer list and management",
    "Review moderation (approve / reject)",
]
for f in admin_features:
    doc.add_paragraph(f, style="List Bullet")

# ── 6. User Guide ──────────────────────────────────────
doc.add_heading("6. User Guide", level=1)
doc.add_heading("6.1 Getting Started", level=2)
doc.add_paragraph(
    "1. Open the website in your browser.\n"
    "2. Browse products on the home page or navigate to /products.\n"
    "3. Click on any product to view details, images, and pricing.\n"
    "4. Use the 'Add to Cart' button to add items to your shopping cart.\n"
    "5. Click the cart icon (top-right) to review cart contents.\n"
    "6. Proceed to checkout to place your order."
)
doc.add_heading("6.2 Creating an Account", level=2)
doc.add_paragraph(
    "Click 'Login' in the top navigation and select 'Sign up'. Enter your name, email, "
    "and password. After signing up, you can log in to access order history, manage your "
    "profile, and save shipping addresses."
)
doc.add_heading("6.3 Placing an Order", level=2)
doc.add_paragraph(
    "1. Add desired products to your cart.\n"
    "2. Click the cart icon and select 'View Cart' or go to /cart.\n"
    "3. Review your items and adjust quantities as needed.\n"
    "4. Click 'Proceed to Checkout'.\n"
    "5. Select or add a shipping address.\n"
    "6. Review your order summary and place the order.\n"
    "7. You will receive a confirmation with an order number for tracking."
)
doc.add_heading("6.4 Tracking an Order", level=2)
doc.add_paragraph(
    "Go to /track and enter your order number to view the current status. You can also "
    "view all your orders on the /orders page after logging in."
)
doc.add_heading("6.5 Managing Your Profile", level=2)
doc.add_paragraph(
    "From the /profile page, you can manage your shipping addresses, update your "
    "information, and view your recent orders."
)

# ── 7. Admin Guide ────────────────────────────────────
doc.add_heading("7. Admin Guide", level=1)
doc.add_heading("7.1 Accessing the Admin Panel", level=2)
doc.add_paragraph(
    "Log in with an admin account and navigate to /admin. The admin panel provides "
    "access to the dashboard, product management, order management, customer list, "
    "and review moderation."
)
doc.add_heading("7.2 Dashboard", level=2)
doc.add_paragraph(
    "The dashboard displays key business metrics including total products, total orders, "
    "total customers, and total revenue. It also shows recent orders for quick reference."
)
doc.add_heading("7.3 Managing Products", level=2)
doc.add_paragraph(
    "From the Products section, you can view all products, add new products, edit "
    "existing ones, or deactivate products. Each product includes name, slug, description, "
    "price, compare-at price (for discounts), material, care instructions, category, and "
    "featured status."
)
doc.add_heading("7.4 Managing Orders", level=2)
doc.add_paragraph(
    "The Orders section lists all customer orders. You can update order status "
    "(pending, confirmed, shipped, delivered, cancelled) and payment status as orders "
    "progress through fulfillment."
)
doc.add_heading("7.5 Review Moderation", level=2)
doc.add_paragraph(
    "Customer reviews require admin approval before appearing on product pages. From "
    "the Reviews section, you can approve or reject reviews."
)

# ── 8. Development & Deployment ────────────────────────
doc.add_heading("8. Development & Deployment", level=1)
doc.add_heading("8.1 Running Locally", level=2)
doc.add_paragraph(
    "Frontend: cd frontend && npm install && npm run dev\n"
    "Backend: cd backend && pip install -r requirements.txt && uvicorn app.main:app --reload\n"
    "The frontend runs on http://localhost:3000 and the backend on http://localhost:8000."
)
doc.add_heading("8.2 Demo Mode", level=2)
doc.add_paragraph(
    "Set NEXT_PUBLIC_USE_DEMO=true in the .env.local file to run the frontend in demo mode. "
    "In demo mode, the application works entirely offline with embedded mock data — no "
    "Supabase connection or backend is required. This is ideal for demonstrations, testing, "
    "and development without external dependencies."
)
doc.add_heading("8.3 Building for Production", level=2)
doc.add_paragraph(
    "Run 'npm run build' in the frontend directory to create an optimized production build. "
    "The output is in the .next/ directory and can be deployed to any Node.js hosting "
    "platform (Vercel, Netlify, Railway, etc.)."
)
doc.add_heading("8.4 Environment Variables", level=2)
env_vars = [
    "NEXT_PUBLIC_USE_DEMO — Set to 'true' to enable demo mode",
    "NEXT_PUBLIC_SUPABASE_URL — Supabase project URL (not needed in demo mode)",
    "NEXT_PUBLIC_SUPABASE_ANON_KEY — Supabase anonymous key (not needed in demo mode)",
    "NEXT_PUBLIC_API_URL — FastAPI backend URL (default: http://localhost:8000)",
]
for var in env_vars:
    doc.add_paragraph(var, style="List Bullet")

# ── 9. Known Limitations & Future Roadmap ──────────────
doc.add_heading("9. Known Limitations & Future Roadmap", level=1)
doc.add_heading("9.1 Current Limitations", level=2)
limitations = [
    "No payment gateway integration (orders are recorded but no real payment processing)",
    "Product images require Supabase storage or local public folder files",
    "Email notifications are not implemented",
    "Search is limited to basic client-side filtering",
    "Inventory/stock tracking is not implemented",
    "Multi-language support is not available",
]
for l in limitations:
    doc.add_paragraph(l, style="List Bullet")

doc.add_heading("9.2 Planned Features", level=2)
planned = [
    "Payment gateway integration (Razorpay/Stripe)",
    "Email order confirmations and shipping notifications",
    "Advanced search with full-text search and filters",
    "Inventory management with low-stock alerts",
    "Size and fit recommendation system",
    "Coupon and discount code system",
    "Customer reviews with image uploads",
    "Wishlist sharing and social features",
    "Multi-language and multi-currency support",
    "Progressive Web App (PWA) support for mobile install",
]
for p in planned:
    doc.add_paragraph(p, style="List Bullet")

# ── 10. Appendix ───────────────────────────────────────
doc.add_heading("10. Appendix", level=1)
doc.add_heading("10.1 Useful Commands", level=2)
commands = [
    ("Start frontend (dev)", "cd frontend && npm run dev"),
    ("Build frontend", "cd frontend && npm run build"),
    ("Lint frontend", "cd frontend && npx next lint"),
    ("Start backend (dev)", "cd backend && uvicorn app.main:app --reload"),
    ("Install backend deps", "cd backend && pip install -r requirements.txt"),
    ("Generate this report", "python generate_report.py"),
]
for label, cmd in commands:
    p = doc.add_paragraph()
    run = p.add_run(f"{label}: ")
    run.bold = True
    run2 = p.add_run(cmd)
    run2.font.name = "Consolas"
    run2.font.size = Pt(9.5)

doc.add_heading("10.2 Key Files Reference", level=2)
files = [
    ("frontend/.env.local", "Environment variables (demo mode toggle)"),
    ("frontend/next.config.ts", "Next.js configuration (images, etc.)"),
    ("frontend/src/lib/config.ts", "Feature flags and demo mode detection"),
    ("frontend/src/lib/demo-data.ts", "Embedded demo product/category data"),
    ("frontend/src/lib/supabase/client.ts", "Supabase client factory with demo fallback"),
    ("frontend/src/lib/context/auth-context.tsx", "Authentication provider"),
    ("frontend/src/lib/stores/cart-store.ts", "Zustand cart state management"),
    ("frontend/src/lib/stores/wishlist-store.ts", "Zustand wishlist state management"),
    ("backend/app/main.py", "FastAPI application entry point"),
    ("backend/app/routes/admin.py", "Admin API routes and validation"),
    ("backend/app/routes/auth.py", "Authentication API routes"),
    ("backend/app/routes/orders.py", "Order management API routes"),
]
for path, desc in files:
    p = doc.add_paragraph()
    run = p.add_run(path)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Consolas"
    p.add_run(f" — {desc}")

# ── Save ────────────────────────────────────────────────
output_path = "D:\\project\\Demo1\\Viraasat_Project_Report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
