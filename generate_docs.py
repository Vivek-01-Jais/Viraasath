from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Styles
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

# ── Title Page ──
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Viraasat — Tech Architecture Overview")
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(128, 0, 32)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("A Complete Guide for Non-Technical Readers")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run("July 2026")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(150, 150, 150)

doc.add_page_break()

# ── Table of Contents ──
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. What is Viraasat?",
    "2. How the Website is Built — Big Picture",
    "3. The Frontend (What You See)",
    "   3.1 Next.js — The Engine Behind the Pages",
    "   3.2 Tailwind CSS — How It Looks",
    "   3.3 TypeScript — Keeping Things Safe",
    "   3.4 Key Frontend Pieces",
    "4. The Backend (What Runs Behind the Scenes)",
    "   4.1 FastAPI — The Brain",
    "   4.2 Supabase — The Memory",
    "   4.3 Razorpay — Taking Payments",
    "   4.4 Key Backend Pieces",
    "5. How Frontend and Backend Talk to Each Other",
    "   5.1 APIs — The Messengers",
    "   5.2 Authentication — Who You Are",
    "   5.3 The Order Flow — End to End",
    "6. Where Everything Lives",
    "7. What Happens When Someone Visits the Site",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ── 1. What is Viraasat ──
doc.add_heading("1. What is Viraasat?", level=1)
doc.add_paragraph(
    "Viraasat is an e-commerce website that sells handcrafted ethnic kurtis and women's wear online. "
    "Customers can browse products by category, search for specific items, add them to a cart, "
    "and pay using credit/debit cards or UPI. Shop owners have a private admin panel to add new "
    "products, manage inventory, view orders, and handle customer reviews."
)
doc.add_paragraph(
    "This document explains the technology used to build Viraasat in simple terms — no coding "
    "knowledge needed."
)

# ── 2. Big Picture ──
doc.add_heading("2. How the Website is Built — Big Picture", level=1)
doc.add_paragraph(
    "Every modern website has two main parts:"
)
doc.add_paragraph(
    "Frontend (the visible part): The pages you see in your browser — the product listing, "
    "the shopping cart, the checkout form. It runs on your phone or computer.",
    style="List Bullet"
)
doc.add_paragraph(
    "Backend (the hidden part): The brain that processes orders, stores product data, "
    "handles payments, and remembers who you are. It runs on a server in the cloud.",
    style="List Bullet"
)
doc.add_paragraph(
    "These two parts talk to each other through APIs (like messengers that carry requests "
    "back and forth). When you click 'Buy Now', your browser asks the backend to place "
    "the order, the backend saves it in the database, and tells your browser 'done'."
)
doc.add_paragraph(
    "Viraasat also uses a third service called Supabase as its database and storage system. "
    "Think of Supabase as a giant filing cabinet where all product details, customer accounts, "
    "and order records are kept."
)

doc.add_page_break()

# ── 3. Frontend ──
doc.add_heading("3. The Frontend (What You See)", level=1)

doc.add_heading("3.1 Next.js — The Engine Behind the Pages", level=2)
doc.add_paragraph(
    "Next.js is a modern framework for building websites. Think of it as a pre-built engine "
    "that handles all the hard parts of making a website fast and reliable:"
)
doc.add_paragraph("It creates pages that load instantly by pre-building them when possible", style="List Bullet")
doc.add_paragraph("It automatically makes images smaller and loads them only when needed", style="List Bullet")
doc.add_paragraph("It handles navigation so users don't have to wait for pages to reload", style="List Bullet")
doc.add_paragraph("It works on all devices — phones, tablets, laptops — automatically", style="List Bullet")
doc.add_paragraph(
    "The current version is Next.js 16, which is the latest stable release."
)

doc.add_heading("3.2 Tailwind CSS — How It Looks", level=2)
doc.add_paragraph(
    "Tailwind CSS is a styling system that controls colors, spacing, fonts, and layout. "
    "Instead of writing custom styles for every element, the developers use ready-made "
    "building blocks. For example, the brand colors used throughout the site:"
)
doc.add_paragraph("Deep maroon (#800020) — used for buttons, headings, and important elements", style="List Bullet")
doc.add_paragraph("Gold (#C5A028 / #B8860B) — used for accents, highlights, and decor", style="List Bullet")
doc.add_paragraph("Cream (#F8F8FF) — used for backgrounds and cards", style="List Bullet")
doc.add_paragraph(
    "The site also supports dark mode — users can switch to a darker color scheme that "
    "is easier on the eyes at night."
)

doc.add_heading("3.3 TypeScript — Keeping Things Safe", level=2)
doc.add_paragraph(
    "TypeScript is JavaScript with extra safety features. It helps catch mistakes before "
    "the website goes live — like making sure a product price is always a number, not text. "
    "The entire frontend is written in TypeScript."
)

doc.add_heading("3.4 Key Frontend Pieces", level=2)
pieces = [
    ("Header & Navigation", "The top bar with the brand name, navigation links (Shop, Track), search bar, cart icon, and login/signup buttons. The 'Shop' button has a dropdown menu showing all product categories."),
    ("Mobile Menu", "A slide-out menu for phone users showing all navigation options including categories, wishlist, track order, sign in, and join."),
    ("Bottom Navigation Bar", "A fixed bar at the bottom of the screen on mobile with icons for Home, Shop, Cart, Wishlist, and Profile."),
    ("Search Command (⌘K)", "A search box that lets users type a product name and see instant results. It opens when pressing Ctrl+K (Windows) or ⌘K (Mac)."),
    ("Product Cards", "The product listing shows items in a grid with image, name, price, discount badge, available sizes, and options to add to wishlist or cart."),
    ("Product Detail Page", "Individual product page with multiple images, size/color selector, price with discount, description, care instructions, reviews, and add-to-cart button."),
    ("Shopping Cart", "A slide-out panel from the right showing all items added to cart with quantity controls and total. Also has a full cart page with order summary."),
    ("Checkout Page", "The payment page where users select shipping address and pay via Razorpay (credit card, UPI, debit card, net banking)."),
    ("Order Tracking", "A page where users can enter their order number to see the status of their delivery."),
    ("Admin Dashboard", "A private area for shop owners to manage products, categories, orders, customers, and reviews. Only accessible by admin accounts."),
    ("Wishlist", "A personal list of saved products that users can come back to later."),
]
for title, desc in pieces:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ── 4. Backend ──
doc.add_heading("4. The Backend (What Runs Behind the Scenes)", level=1)

doc.add_heading("4.1 FastAPI — The Brain", level=2)
doc.add_paragraph(
    "FastAPI is a Python framework that powers the backend server. It handles all business "
    "logic: creating orders, processing payments, verifying admin access, managing addresses, "
    "and handling reviews. It runs on port 8000 and communicates with the frontend through "
    "a set of well-defined API endpoints."
)
doc.add_paragraph(
    "FastAPI was chosen because it is fast (as the name suggests), reliable, and automatically "
    "generates documentation for all the APIs it exposes."
)

doc.add_heading("4.2 Supabase — The Memory", level=2)
doc.add_paragraph(
    "Supabase is an all-in-one backend service that provides three things for Viraasat:"
)
doc.add_paragraph(
    "Database: A PostgreSQL database that stores all information — products, categories, "
    "variants (sizes/colors), images, user profiles, orders, cart items, addresses, reviews, "
    "and wishlists. All data is organized in tables that relate to each other.",
    style="List Bullet"
)
doc.add_paragraph(
    "Storage: An image storage system (like Google Drive) that holds all product photos. "
    "Images are stored in a bucket called 'product-images' and are accessible via public URLs.",
    style="List Bullet"
)
doc.add_paragraph(
    "Authentication: A login system that handles email/password signup and Google login. "
    "It manages user sessions, password resets, and security.",
    style="List Bullet"
)
doc.add_paragraph(
    "The Supabase project is hosted in the cloud at tzdkydvsqqktdjbidzjs.supabase.co."
)

doc.add_heading("4.3 Razorpay — Taking Payments", level=2)
doc.add_paragraph(
    "Razorpay is India's leading payment gateway. When a customer checks out:"
)
doc.add_paragraph("The backend creates a payment order with Razorpay with the exact amount", style="List Bullet")
doc.add_paragraph("The frontend shows a Razorpay payment popup where the customer enters card/UPI details", style="List Bullet")
doc.add_paragraph("Razorpay processes the payment and returns a success or failure response", style="List Bullet")
doc.add_paragraph("The backend verifies the payment signature to prevent fraud", style="List Bullet")
doc.add_paragraph("On success, the order is saved and the cart is emptied", style="List Bullet")
doc.add_paragraph(
    "For testing purposes, Viraasat also has a 'mock payment' mode that works without "
    "real payment keys."
)

doc.add_heading("4.4 Key Backend Pieces", level=2)
backend_pieces = [
    ("Products API", "Manages product listing, searching (full-text search), filtering by category."),
    ("Orders API", "Creates payment orders with Razorpay, places orders, tracks orders by number, lists user orders, and handles Razorpay webhooks (automatic payment status updates)."),
    ("Addresses API", "Manages customer shipping addresses — add, edit, delete, set default."),
    ("Reviews API", "Handles product reviews — customers can submit ratings and comments, admins approve them."),
    ("Admin API", "Full CRUD for products, categories, and orders. Image upload to Supabase Storage. Admin verification and bootstrap (first-user setup). Dashboard statistics (total products, orders, revenue, customers)."),
    ("Auth API", "Handles Google OAuth callback and user profile management."),
]
for title, desc in backend_pieces:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ── 5. How Frontend and Backend Talk ──
doc.add_heading("5. How Frontend and Backend Talk to Each Other", level=1)

doc.add_heading("5.1 APIs — The Messengers", level=2)
doc.add_paragraph(
    "The frontend and backend communicate through HTTP requests — the same technology your "
    "browser uses to load any website. When the frontend needs data, it sends a request to "
    "a specific URL on the backend. For example:"
)
doc.add_paragraph("GET /api/orders/user/123 — 'Give me all orders for user 123'", style="List Bullet")
doc.add_paragraph("POST /api/orders/place — 'Place a new order with these details'", style="List Bullet")
doc.add_paragraph("PUT /api/admin/products/456 — 'Update product 456 with new information'", style="List Bullet")
doc.add_paragraph(
    "The backend runs at http://localhost:8000 during development and responds with data in "
    "JSON format (a structured text format that both computers and humans can read)."
)
doc.add_paragraph(
    "Some data — like product listings and category lists — is fetched directly from Supabase "
    "by the frontend without going through the backend. This makes product pages faster to load."
)

doc.add_heading("5.2 Authentication — Who You Are", level=2)
doc.add_paragraph(
    "When a user logs in (via email or Google), Supabase creates a session token — a digital "
    "pass that proves who the user is. This token is sent with every request to the backend, "
    "so the server knows which user is making the request. The token expires after some time "
    "and the user needs to log in again."
)

doc.add_heading("5.3 The Order Flow — End to End", level=2)
doc.add_paragraph("Here's what happens when a customer buys a product, step by step:")
steps = [
    "Customer browses products on the frontend and clicks 'Add to Cart'",
    "The frontend saves the item in the cart (stored in Supabase)",
    "Customer clicks 'Proceed to Checkout' and selects a shipping address",
    "Frontend asks the backend to create a Razorpay payment order",
    "Backend creates the order with Razorpay and returns payment details",
    "Frontend shows the Razorpay payment popup — customer pays",
    "Razorpay sends a success confirmation to both the frontend and backend",
    "Frontend tells the backend to place the order (with payment proof)",
    "Backend verifies the payment, saves the order, clears the cart",
    "Customer sees 'Order Placed!' and gets an order number to track",
    "The backend webhook receives a final confirmation from Razorpay",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_page_break()

# ── 6. Where Everything Lives ──
doc.add_heading("6. Where Everything Lives", level=1)
doc.add_paragraph(
    "The entire project is stored in a folder called Demo1 on the development computer. "
    "It has two main folders:"
)
doc.add_paragraph(
    "frontend/ — Contains all the website code (Next.js, Tailwind, TypeScript). "
    "This is what users see and interact with.",
    style="List Bullet"
)
doc.add_paragraph(
    "backend/ — Contains all the server code (FastAPI, Python). "
    "This runs behind the scenes and processes orders, payments, etc.",
    style="List Bullet"
)
doc.add_paragraph(
    "The project is deployed using Docker containers — a packaging system that bundles "
    "the frontend and backend with everything they need to run. A single command "
    "(docker compose up) starts the entire website."
)
doc.add_paragraph(
    "Key configuration files:"
)
doc.add_paragraph("frontend/.env.local — Stores API keys and URLs (Supabase, Razorpay, backend URL)", style="List Bullet")
doc.add_paragraph("backend/.env — Stores database credentials, payment keys", style="List Bullet")
doc.add_paragraph("docker-compose.yml — Defines how the frontend and backend containers run together", style="List Bullet")

# ── 7. What Happens When Someone Visits ──
doc.add_heading("7. What Happens When Someone Visits the Site", level=1)
doc.add_paragraph("A complete visitor journey:")
journey = [
    ("Landing on Homepage", "The browser loads the main page. Next.js fetches featured products and categories from Supabase. The page shows a hero banner, category cards, and featured products. All images are optimized — resized and converted to WebP format for fast loading."),
    ("Browsing Products", "User clicks 'Shop' or a category. The products page loads 20 items at a time with pagination. Users can filter by category or search by name/description. Product cards show images, prices, discounts, and available sizes."),
    ("Viewing a Product", "Clicking a product takes the user to a detail page with multiple images, size/color options, full description, material info, care instructions, and customer reviews. The page includes SEO metadata for search engines."),
    ("Adding to Cart", "User selects size and clicks 'Add to Cart'. The item is saved to the cart in Supabase. The cart badge updates in real-time. The slide-out cart panel shows all items with quantity controls."),
    ("Checkout & Payment", "User goes to checkout, selects a saved address, and clicks 'Pay'. A Razorpay popup opens. After successful payment, the order is saved and the cart is cleared. User gets an order number."),
    ("Order Tracking", "User can enter their order number on the Track page to see the current status of their delivery."),
    ("Admin Management", "The shop owner logs into /admin to see the dashboard with total products, orders, revenue, and customers. They can add/edit products, manage categories, fulfill orders, moderate reviews, and upload product images."),
]
for title, desc in journey:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)
    doc.add_paragraph()

# ── Tech Stack Summary Table ──
doc.add_page_break()
doc.add_heading("Technology Stack Summary", level=1)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Layer"
hdr[1].text = "Technology"
hdr[2].text = "Purpose"

tech_stack = [
    ("Frontend Framework", "Next.js 16 (React)", "Builds pages, handles navigation, optimizes images"),
    ("Styling", "Tailwind CSS", "Controls colors, spacing, layout, dark mode"),
    ("Language", "TypeScript", "Adds safety to JavaScript code"),
    ("Animations", "Motion (Framer Motion)", "Adds smooth transitions and animations"),
    ("State Management", "Zustand", "Manages cart, wishlist, and UI state"),
    ("Backend Framework", "FastAPI (Python)", "Handles orders, payments, admin logic"),
    ("Database & Auth", "Supabase (PostgreSQL)", "Stores all data, manages user login"),
    ("Image Storage", "Supabase Storage", "Stores and serves product images"),
    ("Payments", "Razorpay", "Processes customer payments"),
    ("Fonts", "Google Fonts (Inter + Playfair Display)", "Provides the text styles"),
    ("Containerization", "Docker", "Packages the app for deployment"),
]
for layer, tech, purpose in tech_stack:
    row = table.add_row().cells
    row[0].text = layer
    row[1].text = tech
    row[2].text = purpose

doc.add_paragraph()
doc.add_paragraph(
    "This document was generated on " + __import__("datetime").datetime.now().strftime("%B %d, %Y") + "."
)

doc.save("Viraasat_Technology_Overview.docx")
print("Document saved: Viraasat_Technology_Overview.docx")
